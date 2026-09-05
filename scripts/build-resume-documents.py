"""Build the ATS-safe Word resume used for the downloadable PDF.

The document intentionally uses a single text column, native paragraphs,
native heading styles, native lists, and real hyperlinks. Decorative images,
text boxes, and body tables are omitted so reading order survives Word, PDF,
screen readers, and applicant tracking systems.
"""

from pathlib import Path
import shutil

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = ROOT / "output" / "documents" / "ahmed-mahdy-resume.docx"
SITE_DOCX = ROOT / "assets" / "documents" / "ahmed-mahdy-resume.docx"

FONT = "Aptos"
INK = "17212B"
MUTED = "475569"
LINK = "0B5D7A"
RULE = "CBD5E1"


def set_run_font(run, size=None, bold=None, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, *, bold=False, size=11.25):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_node)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_cs)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction, display):
    run = paragraph.add_run()
    set_run_font(run, 8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction_node, separate, value, end))


def add_bottom_border(paragraph, color=RULE, size="6"):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_contact_line(doc):
    paragraph = doc.add_paragraph(style="Contact")
    paragraph.add_run("Cairo, Egypt  |  ")
    add_hyperlink(paragraph, "amahdy59@gmail.com", "mailto:amahdy59@gmail.com", size=11)
    paragraph.add_run("  |  ")
    add_hyperlink(paragraph, "LinkedIn", "https://www.linkedin.com/in/creativemahdy", size=11)
    paragraph.add_run("  |  ")
    add_hyperlink(paragraph, "Portfolio", "https://creativemahdy.space/", size=11)
    for run in paragraph.runs:
        set_run_font(run, 11, color=MUTED)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 1")
    add_bottom_border(paragraph)
    return paragraph


def add_role(doc, title, company, company_url, dates, location, bullets):
    paragraph = doc.add_paragraph(style="Heading 2")
    run = paragraph.add_run(title)
    set_run_font(run, 12, bold=True)
    run = paragraph.add_run("  |  ")
    set_run_font(run, 12, color=MUTED)
    add_hyperlink(paragraph, company, company_url, bold=True, size=12)
    meta = doc.add_paragraph(style="Metadata")
    meta.add_run(f"{dates}  |  {location}  |  Full-time")
    for label, description in bullets:
        bullet = doc.add_paragraph(style="Resume Bullet")
        lead = bullet.add_run(f"{label}: ")
        set_run_font(lead, 11.25, bold=True)
        body = bullet.add_run(description)
        set_run_font(body, 11.25)


def add_project(doc, title, page_url, description, live_label=None, live_url=None):
    paragraph = doc.add_paragraph(style="Project Heading")
    add_hyperlink(paragraph, title, page_url, bold=True, size=11.75)
    if live_label and live_url:
        separator = paragraph.add_run("  |  ")
        set_run_font(separator, 10.5, color=MUTED)
        add_hyperlink(paragraph, live_label, live_url, size=10.5)
    description_paragraph = doc.add_paragraph(description, style="Project Description")
    return description_paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11.25)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = FONT
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("000000")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    title_properties = title._element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    subtitle = doc.styles.add_style("Resume Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = FONT
    subtitle.font.size = Pt(14)
    subtitle.font.bold = True
    subtitle.font.color.rgb = RGBColor.from_string("000000")
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.keep_with_next = True

    contact = doc.styles.add_style("Contact", WD_STYLE_TYPE.PARAGRAPH)
    contact.font.name = FONT
    contact.font.size = Pt(11)
    contact.font.color.rgb = RGBColor.from_string(MUTED)
    contact.paragraph_format.space_after = Pt(12)
    contact.paragraph_format.keep_with_next = True

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = FONT
    heading1.font.size = Pt(13.5)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string("000000")
    heading1.paragraph_format.space_before = Pt(11)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = FONT
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor.from_string("000000")
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(2)
    heading2.paragraph_format.keep_with_next = True

    meta = doc.styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    meta.font.name = FONT
    meta.font.size = Pt(10.25)
    meta.font.bold = True
    meta.font.color.rgb = RGBColor.from_string(MUTED)
    meta.paragraph_format.space_after = Pt(4)
    meta.paragraph_format.keep_with_next = True

    bullet = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet.font.size = Pt(11.25)
    bullet.font.color.rgb = RGBColor.from_string(INK)
    bullet.paragraph_format.left_indent = Inches(0.2)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(3.2)
    bullet.paragraph_format.line_spacing = 1.08

    skill = doc.styles.add_style("Skill Line", WD_STYLE_TYPE.PARAGRAPH)
    skill.font.name = FONT
    skill.font.size = Pt(11.25)
    skill.font.color.rgb = RGBColor.from_string(INK)
    skill.paragraph_format.space_after = Pt(4)

    project_heading = doc.styles.add_style("Project Heading", WD_STYLE_TYPE.PARAGRAPH)
    project_heading.base_style = doc.styles["Heading 2"]
    project_heading.paragraph_format.space_before = Pt(6)
    project_heading.paragraph_format.space_after = Pt(2)

    project_description = doc.styles.add_style("Project Description", WD_STYLE_TYPE.PARAGRAPH)
    project_description.font.name = FONT
    project_description.font.size = Pt(11.25)
    project_description.font.color.rgb = RGBColor.from_string(INK)
    project_description.paragraph_format.space_after = Pt(5)
    project_description.paragraph_format.line_spacing = 1.08

    certification = doc.styles.add_style("Certification", WD_STYLE_TYPE.PARAGRAPH)
    certification.font.name = FONT
    certification.font.size = Pt(11)
    certification.font.color.rgb = RGBColor.from_string(INK)
    certification.paragraph_format.space_after = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    core = doc.core_properties
    core.title = "Ahmed Mahdy Resume"
    core.subject = "UX design and data visualization professional resume"
    core.author = "Ahmed Mahdy"
    core.keywords = "UX design, accessibility, design systems, data visualization, Figma, Power BI, Tableau"
    core.comments = "Single-column ATS-compatible resume"

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    body_language = OxmlElement("w:lang")
    body_language.set(qn("w:val"), "en-US")
    doc.styles["Normal"]._element.rPr.append(body_language)


def configure_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
    name = paragraph.add_run("Ahmed Mahdy Resume  |  Page ")
    set_run_font(name, 8.5, color=MUTED)
    add_field(paragraph, " PAGE ", "1")
    middle = paragraph.add_run(" of ")
    set_run_font(middle, 8.5, color=MUTED)
    add_field(paragraph, " NUMPAGES ", "2")


def build_resume():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    configure_footer(doc)

    doc.add_paragraph("Ahmed Mahdy", style="Title")
    doc.add_paragraph("UX Designer and Data Visualizer", style="Resume Subtitle")
    add_contact_line(doc)

    add_section_heading(doc, "Professional Summary")
    doc.add_paragraph(
        "UX designer and data visualizer with 8+ years of experience creating accessible digital products, enterprise interfaces, learning experiences, and decision-ready dashboards. Combines user research, information architecture, interaction design, prototyping, data analysis, and front-end implementation to move complex work from discovery through validation."
    )

    add_section_heading(doc, "Core Expertise")
    for label, skills in (
        ("UX and product design", "Interaction design, information architecture, user research, wireframing, prototyping, usability testing, design systems, accessibility"),
        ("Data and visualization", "Advanced Excel, Power BI, Tableau, Python, SQL, dashboard design, data storytelling, KPI analysis"),
        ("Tools and delivery", "Figma, Adobe Creative Suite, Microsoft 365, Google Workspace, Looker Studio, Git, GitHub, AI-assisted prototyping"),
    ):
        paragraph = doc.add_paragraph(style="Skill Line")
        lead = paragraph.add_run(f"{label}: ")
        set_run_font(lead, 11.25, bold=True)
        body = paragraph.add_run(skills)
        set_run_font(body, 11.25)

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc,
        "UX Designer",
        "Advansys IS",
        "https://advansys-is.com/",
        "January 2023 - Present",
        "Cairo, Egypt",
        (
            ("Design systems", "Established scalable Figma design systems and reusable interaction patterns for B2B products."),
            ("Enterprise UX", "Led UX and UI design for complex enterprise SaaS platforms, from discovery and information architecture through prototyping and validation."),
            ("AI-assisted delivery", "Used AI design and development tools to accelerate prototyping and research synthesis while retaining design review and accessibility ownership."),
        ),
    )
    add_role(
        doc,
        "Instructional Designer",
        "Schneider Electric",
        "https://www.se.com/eg/en/",
        "July 2018 - January 2023",
        "Cairo, Egypt",
        (
            ("Learning experience design", "Redesigned corporate learning platforms and interfaces to improve engagement and course completion."),
            ("Scalable content", "Built reusable learning assets and templates that streamlined production and supported consistent delivery."),
            ("Learner-centered practice", "Applied learning science, UX principles, and visual storytelling to make complex technical content easier to understand."),
        ),
    )

    doc.add_page_break()

    add_section_heading(doc, "Selected Projects")
    add_project(
        doc,
        "Haj Arafa App",
        "https://creativemahdy.space/en/case-studies/haj-arafa/",
        "Accessible mobile commerce case study centered on search-first navigation, predictable product discovery, and a consolidated guest checkout.",
        "Live project",
        "https://amahdy59.github.io/Hajarafaapp/",
    )
    add_project(
        doc,
        "Cairo International Airport Command Hub",
        "https://creativemahdy.space/en/case-studies/cairo-airport/",
        "Responsive airport operations dashboard that clarifies live status, alerts, and operational priorities through strong information hierarchy and data visualization.",
        "Live project",
        "https://amahdy59.github.io/Cairo-International-Airpot-CIA-Dashboard/",
    )
    add_project(
        doc,
        "HR Management Tool",
        "https://creativemahdy.space/en/case-studies/hr-tool/",
        "Privacy-focused HR SaaS case study for leave requests, role-based approvals, and clear request-status visibility across desktop and mobile.",
        "Live project",
        "https://amahdy59.github.io/hr-tool-v2/",
    )
    add_project(
        doc,
        "Azkar App Daily Fortress",
        "https://creativemahdy.space/en/case-studies/azkar-app/",
        "Arabic-first remembrance and prayer-time application with accessible reading controls, offline support, calm content presentation, and clear progress.",
        "Live project",
        "https://amahdy59.github.io/Azkarapp/#/home",
    )
    add_project(
        doc,
        "Data-Driven LEGO Explorer",
        "https://creativemahdy.space/en/case-studies/lego-explorer/",
        "Power BI dashboard supported by Python analysis for comparing LEGO themes, age ranges, pricing, and piece counts through interactive visual exploration.",
        "Dashboard",
        "https://mavenshowcase.com/project/24967",
    )

    add_section_heading(doc, "Education")
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.add_run("Diploma in Education and Instructional Technology")
    meta = doc.add_paragraph(style="Metadata")
    add_hyperlink(meta, "Information Technology Institute", "https://iti.gov.eg/iti/home", bold=True, size=10.25)
    meta.add_run("  |  September 2016 - June 2017")
    doc.add_paragraph("Focused on instructional design, educational technology, digital learning production, and learner-centered content.", style="Project Description")
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.add_run("Bachelor's Degree in Radio and Television")
    meta = doc.add_paragraph(style="Metadata")
    add_hyperlink(meta, "Minufiya University", "https://www.menofia.edu.eg/", bold=True, size=10.25)
    meta.add_run("  |  September 2009 - June 2013")
    doc.add_paragraph("Studied broadcast communication, media production, and visual storytelling.", style="Project Description")

    add_section_heading(doc, "Certifications")
    certifications = (
        ("Google Data Analytics Professional Certificate", "Google, 2024", "https://www.coursera.org/account/accomplishments/professional-cert/0OYUGNC1DRF5"),
        ("Tableau Business Intelligence Analyst", "Tableau, 2024", "https://www.coursera.org/account/accomplishments/professional-cert/JQPVYZ6VYF52"),
        ("Excel Skills for Data Analytics and Visualization", "Macquarie University, 2024", "https://www.coursera.org/account/accomplishments/specialization/5XR5AHBP79KW"),
        ("Excel Skills for Business", "Macquarie University, 2024", "https://www.coursera.org/account/accomplishments/specialization/DCL8H4YT5UWG"),
        ("Google UX Design Professional Certificate", "Google, 2023", "https://www.coursera.org/account/accomplishments/professional-cert/49PCMPYYGLJV"),
    )
    for title, issuer, url in certifications:
        paragraph = doc.add_paragraph(style="Certification")
        add_hyperlink(paragraph, title, url, bold=True, size=11)
        suffix = paragraph.add_run(f"  |  {issuer}")
        set_run_font(suffix, 11, color=MUTED)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    SITE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    shutil.copy2(OUTPUT_DOCX, SITE_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_resume()
